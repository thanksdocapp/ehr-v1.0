#!/usr/bin/env python3
"""Generate ThankDoc EHR Client Proposal as DOCX from markdown content."""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

def add_heading(doc, text, level=1):
    """Add a heading with appropriate style."""
    if level == 1:
        p = doc.add_heading(text, level=0)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        doc.add_heading(text, level=level)

def add_paragraph(doc, text):
    """Add a paragraph."""
    if text.strip():
        doc.add_paragraph(text.strip())

def add_bullet_list(doc, items):
    """Add a bullet list."""
    for item in items:
        p = doc.add_paragraph(item.strip(), style='List Bullet')

def add_table_from_rows(doc, headers, rows):
    """Add a table with headers and rows."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
    for row_idx, row in enumerate(rows):
        row_cells = table.rows[row_idx + 1].cells
        for i, cell in enumerate(row):
            if i < len(row_cells):
                row_cells[i].text = str(cell)
    doc.add_paragraph()

def main():
    doc = Document()
    
    # Title
    title = doc.add_paragraph()
    title_run = title.add_run('ThankDoc EHR — Client Proposal')
    title_run.bold = True
    title_run.font.size = Pt(22)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph()
    subtitle_run = subtitle.add_run('A modern, compliant Electronic Health Record system for clinics and multi-doctor practices')
    subtitle_run.italic = True
    subtitle_run.font.size = Pt(12)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    
    # Executive Summary
    add_heading(doc, 'Executive Summary', 1)
    add_paragraph(doc, 
        'ThankDoc EHR is a full-featured hospital and clinic management system designed to streamline operations, '
        'improve patient care, and support growth. Built on secure, enterprise-grade technology, it combines '
        'clinical documentation, appointment booking, billing, and analytics in one integrated platform—with '
        'built-in support for multi-doctor practices, multiple booking links, and marketing attribution.')
    doc.add_paragraph()
    
    # Why ThankDoc EHR?
    add_heading(doc, 'Why ThankDoc EHR?', 1)
    
    add_heading(doc, 'For Practice Owners & Administrators', 2)
    add_bullet_list(doc, [
        'Single platform for patients, appointments, records, billing, and reporting',
        'Role-based access so staff see only what they need (Admin, Doctor, Nurse, Receptionist)',
        'Department and clinic structure for multi-site or multi-specialty practices',
        'Revenue and performance analytics to support business decisions',
    ])
    
    add_heading(doc, 'For Clinicians', 2)
    add_bullet_list(doc, [
        'Structured clinical notes (PC, HPC, PMH, DH, ICE, Plan) with templates',
        'Vital signs with automatic BMI calculation',
        'Document attachments linked to consultations',
        'Prescriptions and lab reports integrated with medical records',
        'Pre-consultation verification checklist for compliance',
    ])
    
    add_heading(doc, 'For Patients', 2)
    add_bullet_list(doc, [
        'Online booking with doctor-, clinic-, or service-specific links',
        'Patient portal for appointments and records',
        'Email and SMS reminders to reduce no-shows',
        'Secure access to their health information',
    ])
    
    add_heading(doc, 'For Marketing & Growth', 2)
    add_bullet_list(doc, [
        'Unique booking links per doctor, clinic, or service',
        'Conversion tracking for Google Ads and other campaigns',
        'Multi-agency attribution via UTM parameters (utm_source, utm_medium, utm_campaign)',
        'Iframe embedding for integration with practice websites (e.g. WordPress)',
    ])
    doc.add_paragraph()
    
    # Core Capabilities
    add_heading(doc, 'Core Capabilities', 1)
    add_table_from_rows(doc, 
        ['Module', 'Features'],
        [
            ['Patients', 'Registration, medical history, insurance, portal access'],
            ['Appointments', 'Online booking, calendar, reminders, status tracking'],
            ['Medical Records', 'EHR, diagnoses, treatments, attachments, audit trail'],
            ['Prescriptions', 'Digital prescriptions, drug history, interaction checks'],
            ['Lab Reports', 'Test ordering, results, report generation'],
            ['Billing', 'Invoicing, payments, online payment gateways'],
            ['Notifications', 'Email, SMS, appointment reminders, clinical alerts'],
            ['Reporting', 'Dashboards, analytics, custom reports, exports'],
        ])
    
    # Features
    add_heading(doc, 'Features', 1)
    
    features = [
        ('Staff & User Management', [
            'Complete staff registration and profile management',
            'Role-based access control (Admin, Doctor, Nurse, Receptionist, etc.)',
            'Department and clinic structure for multi-site practices',
            'Two-factor authentication (2FA) for enhanced security',
            'Custom role menu visibility and permissions',
            'User activity logging and audit trails',
        ]),
        ('Patient Management', [
            'Patient registration with comprehensive profiles',
            'Medical history, allergies, and drug history tracking',
            'Patient portal for self-service access',
            'Insurance and billing information management',
            'Patient alerts and clinical flags',
            'Patient search and filtering',
            'GP/carer consent and sharing preferences',
        ]),
        ('Appointment System', [
            'Online appointment booking (public and staff)',
            'Doctor-, clinic-, and service-specific booking links',
            'Calendar view with availability management',
            'Automated email and SMS notifications',
            'Appointment status tracking (scheduled, confirmed, completed, cancelled)',
            'Reschedule and cancel from patient dashboard',
            'Video consultation links (Zoom, Google Meet, Whereby, etc.)',
            'Follow-up reminders',
        ]),
        ('Medical Records (EHR)', [
            'Structured clinical notes (PC, HPC, PMH, DH, ICE, Plan)',
            'Diagnosis and treatment tracking',
            'Vital signs with automatic BMI calculation',
            'Document and file attachments linked to consultations',
            'Pre-consultation verification checklist',
            'Copy-forward from previous records',
            'Private records option for sensitive consultations',
            'Lab results and prescriptions integration',
            'Follow-up date tracking',
            'CSV import for bulk records',
        ]),
        ('Prescriptions', [
            'Digital prescription management',
            'Link prescriptions to medical records',
            'Drug history and allergy checks',
            'Prescription templates',
            'Print and export',
        ]),
        ('Laboratory Management', [
            'Lab test ordering and tracking',
            'Results management and report generation',
            'Link lab reports to medical records and patients',
            'Quality control tracking',
        ]),
        ('Billing & Finance', [
            'Automated billing and invoice generation',
            'Multiple payment gateways (Stripe, PayPal, Paystack, Flutterwave, BTC Pay)',
            'Payment tracking and transaction history',
            'Online payment for patients',
            'Public payment links for services',
            'Financial reporting',
        ]),
        ('Documents & Forms', [
            'Document templates (letters, forms, consultation forms)',
            'Patient documents with e-signature support',
            'Document delivery and tracking (email open/click)',
            'Form requests and submissions',
            'Generated documents linked to patients',
        ]),
        ('Communication', [
            'Email templates (customisable)',
            'SMS templates and sending',
            'In-app notifications',
            'Patient–doctor messaging (GP email)',
            'Contact form and message management',
        ]),
        ('Marketing & Booking', [
            'Unique booking links per doctor, clinic, or service',
            'Google Tag Manager (GTM) integration',
            'Conversion tracking for Google Ads',
            'Multi-agency attribution via UTM parameters',
            'Iframe embedding for practice websites (e.g. WordPress)',
        ]),
        ('Website & Frontend', [
            'Customisable homepage and sections',
            'Banner slides and key features',
            'Testimonials and about stats',
            'FAQ management',
            'SEO settings (meta tags, Google Analytics, GTM)',
            'Dynamic theme and branding',
            'Mobile-responsive design',
        ]),
        ('Integrations & API', [
            'RESTful API for mobile apps and integrations',
            'Payment gateway webhooks',
            'Third-party module support',
            'Document tracking (open/click)',
        ]),
        ('Reporting & Analytics', [
            'Comprehensive dashboards (admin, staff, patient)',
            'Revenue and patient analytics',
            'Doctor and department performance metrics',
            'Advanced reports and audit trail',
            'Consultations report',
            'Export to CSV/Excel',
        ]),
        ('Security & Compliance', [
            'GDPR-aware design with consent controls',
            'Encrypted clinical data (PHI)',
            'Audit logging for access and changes',
            'Role-based permissions',
            'Private records option',
        ]),
    ]
    
    for section_title, items in features:
        add_heading(doc, section_title, 2)
        add_bullet_list(doc, items)
        doc.add_paragraph()
    
    # Compliance & Security
    add_heading(doc, 'Compliance & Security', 1)
    add_bullet_list(doc, [
        'GDPR-aware design with consent and data handling controls',
        'Encrypted clinical data for sensitive information',
        'Audit logging for access and changes',
        'Role-based permissions to limit data access',
        'Private records option for sensitive consultations',
    ])
    doc.add_paragraph()
    
    # Technical Highlights
    add_heading(doc, 'Technical Highlights', 1)
    add_bullet_list(doc, [
        'Modern stack: Laravel (PHP), Bootstrap, responsive design',
        'RESTful API for integrations and mobile apps',
        'Iframe-ready booking for embedding on practice websites',
        'Configurable branding, themes, and settings',
        'Scalable architecture for growing practices',
    ])
    doc.add_paragraph()
    
    # Implementation & Support
    add_heading(doc, 'Implementation & Support', 1)
    add_bullet_list(doc, [
        'Standard installation with guided setup',
        'Data migration support for existing systems',
        'Training for administrators and clinical staff',
        'Documentation and configuration guides',
        'Ongoing updates for security and new features',
    ])
    doc.add_paragraph()
    
    # Next Steps
    add_heading(doc, 'Next Steps', 1)
    add_bullet_list(doc, [
        'Discovery call — Discuss your practice size, workflows, and priorities',
        'Demo — Walkthrough of the system tailored to your use case',
        'Proposal — Scope, timeline, and investment',
        'Pilot — Trial period with a subset of users or departments',
        'Go-live — Full deployment with training and support',
    ])
    doc.add_paragraph()
    
    # Footer
    footer = doc.add_paragraph()
    footer_run = footer.add_run('ThankDoc EHR — Built for modern healthcare practices.')
    footer_run.italic = True
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Save
    output_path = 'c:\\Users\\chukw\\Documents\\ehr-v1.0\\docs\\ThankDoc_EHR_Client_Proposal.docx'
    doc.save(output_path)
    print(f'Saved: {output_path}')

if __name__ == '__main__':
    main()
